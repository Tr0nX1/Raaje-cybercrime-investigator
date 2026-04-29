import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class NoticeGenerator:
    """
    Handles file creation in .txt and .docx formats with styling and letterhead support.
    """
    
    def __init__(self, output_dir, letterhead_path=None, style_config=None):
        self.output_dir = output_dir
        self.letterhead_path = letterhead_path
        
        # Default Style Configuration
        self.style_config = {
            'alignment': 'CENTER',
            'font_name': 'Arial',
            'font_size': 11,
            'letterhead_width': 6.0, # Inches
            'margins': 1.0 # Inches (not implemented yet in this basic version but placeholder)
        }
        if style_config:
            self.style_config.update(style_config)

        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def generate_txt(self, filename, content):
        """Generates a simple .txt file."""
        if not filename.endswith('.txt'):
            filename += '.txt'
        
        path = os.path.join(self.output_dir, filename)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return path

    def _get_alignment(self, align_str):
        """Maps string to WD_ALIGN_PARAGRAPH enum."""
        mapping = {
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return mapping.get(align_str.upper(), WD_ALIGN_PARAGRAPH.CENTER)

    def generate_docx(self, filename, content):
        """
        Generates a .docx file with letterhead and styling.
        """
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        path = os.path.join(self.output_dir, filename)
        doc = Document()
        
        # 1. Letterhead Handling
        if self.letterhead_path and os.path.exists(self.letterhead_path):
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = self._get_alignment(self.style_config['alignment'])
            
            run = header_para.add_run()
            run.add_picture(self.letterhead_path, width=Inches(self.style_config['letterhead_width']))

        # 2. Global Styling (Font)
        style = doc.styles['Normal']
        style.font.name = self.style_config['font_name']
        style.font.size = Pt(self.style_config['font_size'])

        # 3. Content Addition
        # Split content by newlines and add to document
        for line in content.split('\n'):
            p = doc.add_paragraph(line)
            # You could also apply alignment to body text if needed, 
            # but usually it's JUSTIFY or LEFT.
            
        doc.save(path)
        return path

    def create_filename(self, bank_name):
        """Creates a sanitized filename."""
        sanitized = "".join([c if c.isalnum() else "_" for c in bank_name])
        date_str = datetime.now().strftime("%Y%m%d_%H%M")
        return f"Notice_{sanitized}_{date_str}"
